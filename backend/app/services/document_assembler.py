"""Document Assembler — builds the final DOCX from generated section content.

Strategy:
  1. Clone the template DOCX.
  2. Clear all body content below the cover page (if any), preserving styles.
  3. For each generated section, insert:
     - Heading paragraph with the correct heading style
     - Content elements: narrative paragraphs, bullet lists, tables with captions
  4. Apply template-matched formatting to all inserted elements:
     - Font family, size, color from style catalog
     - Table borders, header colors, alternating row colors
  5. Build a visible Table of Contents from headings (if requested).
  6. Preserve headers, footers, and page layout from the template.
"""
from __future__ import annotations

import copy
import logging
import shutil
from pathlib import Path
from typing import Any

log = logging.getLogger(__name__)


class DocumentAssembler:
    """Assembles a professional DOCX from generated section content."""

    def assemble(
        self,
        template_path: Path,
        target_path: Path,
        generated_sections: list[dict],
        template_analysis: dict,
    ) -> tuple[bool, list[str]]:
        """Assemble the final document.

        Args:
            template_path: Source DOCX template.
            target_path: Where to write the assembled document.
            generated_sections: Output of SectionGenerator.generate_all_sections().
            template_analysis: Output of TemplateAnalyzer.analyze().

        Returns:
            (success, list_of_section_summaries)
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.oxml.ns import qn
        except ImportError:
            raise RuntimeError("python-docx is required.")

        # Clone template
        shutil.copy2(template_path, target_path)

        doc = Document(str(target_path))

        # Clear existing body content (keep styles intact)
        self._clear_body(doc, template_analysis)

        style_catalog = template_analysis.get("style_catalog", {})
        table_style = template_analysis.get("table_style", {})

        summaries: list[str] = []
        toc_headings: list[dict] = []

        for section in generated_sections:
            heading_text = section.get("heading", "")
            heading_level = section.get("heading_level", 1)
            style_name = section.get("style_name", f"Heading {heading_level}")
            elements = section.get("elements", [])

            if not heading_text:
                continue

            # Insert heading
            heading_para = self._add_heading(doc, heading_text, style_name)
            toc_headings.append({"text": heading_text, "level": heading_level})

            # Insert elements
            has_content = False
            for element in elements:
                el_type = element.get("type", "paragraph")

                if el_type == "toc_placeholder":
                    self._add_toc(doc, toc_headings)
                    has_content = True

                elif el_type == "paragraph":
                    text = element.get("text", "")
                    if text:
                        self._add_paragraph(doc, text, element.get("style", "Normal"), style_catalog)
                        has_content = True

                elif el_type in ("bullet_list", "numbered_list"):
                    items = element.get("items", [])
                    list_style = "List Bullet" if el_type == "bullet_list" else "List Number"
                    effective_style = element.get("style", list_style)
                    for item in items:
                        if item:
                            self._add_paragraph(doc, item, effective_style, style_catalog)
                    if items:
                        has_content = True

                elif el_type == "table_caption":
                    text = element.get("text", "")
                    if text:
                        self._add_paragraph(doc, text, element.get("style", "Caption"), style_catalog)

                elif el_type == "table":
                    headers = element.get("headers", [])
                    rows = element.get("rows", [])
                    if headers:
                        self._add_table(doc, headers, rows, table_style)
                        has_content = True

            if has_content:
                summaries.append(f"Generated section: {heading_text}")

        # Add page break between major sections for readability
        # (already done by heading styles in most templates)

        doc.save(str(target_path))
        log.info("DocumentAssembler: saved %s with %d sections", target_path, len(summaries))
        return True, summaries

    # ------------------------------------------------------------------
    # Body Clearing
    # ------------------------------------------------------------------

    def _clear_body(self, doc, template_analysis: dict) -> None:
        """Remove all body content while preserving styles.

        If the template has a cover page (detected via 'has_cover_page'),
        we keep the cover page content and clear only after the first
        section break / page break.
        """
        from docx.oxml.ns import qn

        body = doc.element.body
        has_cover = template_analysis.get("has_cover_page", False)

        # Collect all body children (paragraphs, tables, sectPr)
        children = list(body)

        # Find sectPr (section properties) — must be kept
        sectPr = body.find(qn("w:sectPr"))

        # Determine where to start clearing
        # If cover page detected, keep first paragraph cluster before first heading
        skip_until = 0
        if has_cover:
            for i, child in enumerate(children):
                tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                if tag == "p":
                    style_elem = child.find(f".//{qn('w:pStyle')}")
                    if style_elem is not None:
                        style_val = style_elem.get(qn("w:val"), "")
                        if "Heading" in style_val:
                            skip_until = i
                            break

        # Remove all non-cover-page children
        for child in children[skip_until:]:
            if child is sectPr:
                continue
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag in ("p", "tbl", "sdt"):
                body.remove(child)

        # Ensure sectPr stays at end
        if sectPr is not None:
            body.remove(sectPr)
            body.append(sectPr)

    # ------------------------------------------------------------------
    # Element Insertion
    # ------------------------------------------------------------------

    def _add_heading(self, doc, text: str, style_name: str):
        """Add a heading paragraph, falling back to a numbered heading if style missing."""
        from docx.shared import Pt
        try:
            # Try exact style name first
            para = doc.add_paragraph(style=style_name)
        except Exception:
            # Fallback: try "Heading 1" etc.
            level = 1
            try:
                level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                pass
            try:
                para = doc.add_paragraph(style=f"Heading {level}")
            except Exception:
                para = doc.add_paragraph()

        run = para.add_run(text)
        return para

    def _add_paragraph(self, doc, text: str, style_name: str, style_catalog: dict):
        """Add a body paragraph with template-matched formatting."""
        try:
            para = doc.add_paragraph(style=style_name)
        except Exception:
            para = doc.add_paragraph()

        run = para.add_run(text)

        # Apply additional formatting from style catalog if style has explicit settings
        # (python-docx paragraph styles already carry most formatting, so we only
        # override where the catalog explicitly specifies something not captured by style)
        style_info = style_catalog.get(style_name, {})
        self._apply_run_formatting(run, style_info)

        return para

    def _apply_run_formatting(self, run, style_info: dict) -> None:
        """Apply explicit run-level formatting from style catalog."""
        from docx.shared import Pt, RGBColor
        try:
            if style_info.get("font_family"):
                run.font.name = style_info["font_family"]
            if style_info.get("font_size_pt"):
                run.font.size = Pt(style_info["font_size_pt"])
            if style_info.get("bold") is not None:
                run.font.bold = style_info["bold"]
            if style_info.get("italic") is not None:
                run.font.italic = style_info["italic"]
            color_hex = style_info.get("color_hex")
            if color_hex and len(color_hex) == 6:
                r = int(color_hex[0:2], 16)
                g = int(color_hex[2:4], 16)
                b = int(color_hex[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
        except Exception as exc:
            log.debug("Run formatting failed: %s", exc)

    def _add_table(
        self,
        doc,
        headers: list[str],
        rows: list[list[str]],
        table_style: dict,
    ):
        """Add a professionally formatted table matching the template style."""
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        num_cols = len(headers)
        num_rows = len(rows) + 1  # +1 for header row

        # Try to use the template's table style
        style_name = table_style.get("style_name", "Table Grid")
        try:
            table = doc.add_table(rows=num_rows, cols=num_cols, style=style_name)
        except Exception:
            try:
                table = doc.add_table(rows=num_rows, cols=num_cols, style="Table Grid")
            except Exception:
                table = doc.add_table(rows=num_rows, cols=num_cols)

        table.autofit = True

        # Set header row
        header_row = table.rows[0]
        header_bg = table_style.get("header_bg_hex")
        header_fc = table_style.get("header_font_color_hex")
        header_bold = table_style.get("header_bold", True)

        for col_idx, header_text in enumerate(headers):
            cell = header_row.cells[col_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(header_text)
            run.font.bold = header_bold

            if header_fc and len(header_fc) == 6:
                try:
                    r = int(header_fc[0:2], 16)
                    g = int(header_fc[2:4], 16)
                    b = int(header_fc[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)
                except Exception:
                    pass

            if header_bg and len(header_bg) == 6:
                self._set_cell_bg(cell, header_bg)

        # Data rows
        alt_colors = table_style.get("alt_row_colors", [])
        for row_idx, row_data in enumerate(rows):
            table_row = table.rows[row_idx + 1]
            for col_idx, cell_text in enumerate(row_data):
                if col_idx >= num_cols:
                    break
                cell = table_row.cells[col_idx]
                cell.text = str(cell_text)

            # Alternate row coloring
            if alt_colors and len(alt_colors) >= 1:
                bg = alt_colors[row_idx % len(alt_colors)]
                if bg and len(bg) == 6:
                    for cell in table_row.cells:
                        self._set_cell_bg(cell, bg)

        return table

    def _set_cell_bg(self, cell, hex_color: str) -> None:
        """Set the background fill color of a table cell."""
        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), hex_color.upper())
            # Remove existing shd if any
            existing = tcPr.find(qn("w:shd"))
            if existing is not None:
                tcPr.remove(existing)
            tcPr.append(shd)
        except Exception as exc:
            log.debug("Cell background set failed: %s", exc)

    def _add_toc(self, doc, headings: list[dict]) -> None:
        """Insert a native Word Table of Contents field."""
        from docx.shared import Pt
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        # Add TOC heading paragraph
        try:
            toc_para = doc.add_paragraph("Table of Contents", style="TOC Heading")
        except Exception:
            toc_para = doc.add_paragraph("Table of Contents")
            for run in toc_para.runs:
                run.font.bold = True
                run.font.size = Pt(14)

        fld_xml = (
            r'<w:p %s>'
            r'  <w:r>'
            r'    <w:fldChar w:fldCharType="begin"/>'
            r'  </w:r>'
            r'  <w:r>'
            r'    <w:instrText xml:space="preserve"> TOC \o "1-3" \h \z \u </w:instrText>'
            r'  </w:r>'
            r'  <w:r>'
            r'    <w:fldChar w:fldCharType="separate"/>'
            r'  </w:r>'
            r'  <w:r>'
            r'    <w:fldChar w:fldCharType="end"/>'
            r'  </w:r>'
            r'</w:p>' % nsdecls('w')
        )
        p_toc = parse_xml(fld_xml)
        doc.element.body.append(p_toc)

        from app.services.document_processor import _enable_update_fields
        _enable_update_fields(doc)
