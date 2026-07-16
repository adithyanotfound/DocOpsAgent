"""Knowledge Base Document Processor.

Parses uploaded documents (PDF, DOCX, TXT, MD) into text chunks suitable
for embedding and semantic retrieval.

Chunking strategy:
  - Target ~500 tokens per chunk (~350 words) with 50-token overlap.
  - Preserve section headings in chunk metadata for citation context.
  - PDF: page-by-page extraction, then paragraph merging.
  - DOCX: heading-aware extraction preserving document hierarchy.
  - TXT/MD: Paragraph-based splitting on double newlines.
"""
from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any

log = logging.getLogger(__name__)

# Rough heuristic: ~1.3 chars per token for English prose
CHARS_PER_TOKEN = 1.3
TARGET_CHUNK_TOKENS = 500
OVERLAP_TOKENS = 50
TARGET_CHUNK_CHARS = int(TARGET_CHUNK_TOKENS * CHARS_PER_TOKEN)
OVERLAP_CHARS = int(OVERLAP_TOKENS * CHARS_PER_TOKEN)


class KBProcessor:
    """Parses and chunks documents for the knowledge base."""

    SUPPORTED_TYPES = {"pdf", "docx", "txt", "md"}

    def process_document(self, file_path: Path, file_type: str) -> list[dict]:
        """Parse a document and return a list of text chunks with metadata.

        Each chunk dict has:
          - text: the chunk content
          - chunk_index: 0-based position
          - metadata: source info (page, section, filename, etc.)
        """
        file_type = file_type.lower().lstrip(".")
        if file_type not in self.SUPPORTED_TYPES:
            raise ValueError(f"Unsupported file type: {file_type}")

        filename = file_path.name

        if file_type == "pdf":
            sections = self._parse_pdf(file_path, filename)
        elif file_type == "docx":
            sections = self._parse_docx(file_path, filename)
        else:  # txt or md
            sections = self._parse_text(file_path, filename)

        chunks = self._chunk_sections(sections)
        return chunks

    # ------------------------------------------------------------------
    # Parsers
    # ------------------------------------------------------------------

    def _parse_pdf(self, path: Path, filename: str) -> list[dict]:
        """Extract text page-by-page from PDF."""
        try:
            import PyPDF2
        except ImportError:
            # Fallback: attempt raw text read
            log.warning("PyPDF2 not installed; trying plain text fallback for PDF.")
            return self._parse_text(path, filename)

        sections: list[dict] = []
        try:
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page_num, page in enumerate(reader.pages, start=1):
                    text = page.extract_text() or ""
                    text = text.strip()
                    if not text:
                        continue
                    # Split page text into paragraphs
                    paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
                    for para in paragraphs:
                        sections.append({
                            "text": para,
                            "metadata": {
                                "source": filename,
                                "page": page_num,
                                "type": "pdf_paragraph",
                            },
                        })
        except Exception as exc:
            log.error("PDF parsing failed for %s: %s", path, exc)
            raise

        return sections

    def _parse_docx(self, path: Path, filename: str) -> list[dict]:
        """Extract paragraphs from DOCX, preserving heading hierarchy."""
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError("python-docx is required but not installed.")

        sections: list[dict] = []
        current_heading: str | None = None
        current_heading_level: int | None = None

        try:
            doc = Document(path)
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                style_name = para.style.name if para.style else "Normal"
                is_heading = style_name.startswith("Heading")
                heading_level = None

                if is_heading:
                    try:
                        heading_level = int(style_name.split()[-1])
                    except (ValueError, IndexError):
                        heading_level = 1
                    current_heading = text
                    current_heading_level = heading_level
                    # Add heading itself as a section marker
                    sections.append({
                        "text": text,
                        "metadata": {
                            "source": filename,
                            "section": text,
                            "heading_level": heading_level,
                            "type": "heading",
                        },
                    })
                else:
                    sections.append({
                        "text": text,
                        "metadata": {
                            "source": filename,
                            "section": current_heading,
                            "heading_level": current_heading_level,
                            "type": "paragraph",
                        },
                    })

            # Also extract table content
            for table_idx, table in enumerate(doc.tables, start=1):
                rows_text: list[str] = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    row_text = " | ".join(c for c in cells if c)
                    if row_text:
                        rows_text.append(row_text)
                if rows_text:
                    table_text = "\n".join(rows_text)
                    sections.append({
                        "text": table_text,
                        "metadata": {
                            "source": filename,
                            "section": current_heading,
                            "type": "table",
                            "table_index": table_idx,
                        },
                    })
        except Exception as exc:
            log.error("DOCX KB parsing failed for %s: %s", path, exc)
            raise

        return sections

    def _parse_text(self, path: Path, filename: str) -> list[dict]:
        """Parse plain text or markdown files."""
        try:
            text = path.read_text(encoding="utf-8", errors="replace")
        except Exception as exc:
            log.error("Text parsing failed for %s: %s", path, exc)
            raise

        # Split on double newlines (markdown paragraphs) or single newlines for long lines
        paragraphs = re.split(r"\n{2,}", text)
        sections: list[dict] = []
        current_heading: str | None = None

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # Detect markdown headings
            heading_match = re.match(r"^(#{1,6})\s+(.+)$", para, re.MULTILINE)
            if heading_match:
                current_heading = heading_match.group(2).strip()
                level = len(heading_match.group(1))
                sections.append({
                    "text": current_heading,
                    "metadata": {
                        "source": filename,
                        "section": current_heading,
                        "heading_level": level,
                        "type": "heading",
                    },
                })
            else:
                sections.append({
                    "text": para,
                    "metadata": {
                        "source": filename,
                        "section": current_heading,
                        "type": "paragraph",
                    },
                })

        return sections

    # ------------------------------------------------------------------
    # Chunking
    # ------------------------------------------------------------------

    def _chunk_sections(self, sections: list[dict]) -> list[dict]:
        """Combine small sections into chunks and split large sections.

        Attempts to keep section context intact. Each chunk includes
        its source metadata from the first constituent section.
        """
        chunks: list[dict] = []
        current_text = ""
        current_meta: dict[str, Any] = {}

        for section in sections:
            text = section["text"]
            meta = section["metadata"]

            # If current text is empty, start a new chunk
            if not current_text:
                current_text = text
                current_meta = meta
                continue

            # If adding this section would exceed the target, flush first
            if len(current_text) + len(text) + 1 > TARGET_CHUNK_CHARS:
                # Flush current
                if current_text.strip():
                    chunks.append({
                        "text": current_text.strip(),
                        "chunk_index": len(chunks),
                        "metadata": current_meta,
                    })
                    # Carry overlap from end of current chunk
                    overlap_start = max(0, len(current_text) - OVERLAP_CHARS)
                    current_text = current_text[overlap_start:].strip() + "\n" + text
                    current_meta = meta
                else:
                    current_text = text
                    current_meta = meta
            else:
                current_text += "\n" + text

        # Flush remaining
        if current_text.strip():
            chunks.append({
                "text": current_text.strip(),
                "chunk_index": len(chunks),
                "metadata": current_meta,
            })

        # Handle individual oversized sections by splitting them
        final_chunks: list[dict] = []
        for chunk in chunks:
            if len(chunk["text"]) > TARGET_CHUNK_CHARS * 2:
                sub = self._split_large_text(chunk["text"], chunk["metadata"])
                for i, s in enumerate(sub):
                    s["chunk_index"] = len(final_chunks)
                    final_chunks.append(s)
            else:
                chunk["chunk_index"] = len(final_chunks)
                final_chunks.append(chunk)

        return final_chunks

    def _split_large_text(self, text: str, meta: dict) -> list[dict]:
        """Split a large text block into overlapping chunks at sentence boundaries."""
        # Try to split at sentence boundaries
        sentences = re.split(r"(?<=[.!?])\s+", text)
        chunks: list[dict] = []
        current = ""

        for sentence in sentences:
            if len(current) + len(sentence) > TARGET_CHUNK_CHARS:
                if current.strip():
                    chunks.append({
                        "text": current.strip(),
                        "chunk_index": 0,
                        "metadata": meta,
                    })
                    # Start new chunk with overlap
                    overlap_start = max(0, len(current) - OVERLAP_CHARS)
                    current = current[overlap_start:].strip() + " " + sentence
                else:
                    current = sentence
            else:
                current += " " + sentence

        if current.strip():
            chunks.append({
                "text": current.strip(),
                "chunk_index": 0,
                "metadata": meta,
            })

        return chunks
