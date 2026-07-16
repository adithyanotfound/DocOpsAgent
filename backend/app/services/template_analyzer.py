"""Template Analyzer — deep extraction of DOCX template structure and styling.

Extracts:
  1. Structural skeleton: ordered sections with heading levels and names
  2. Style catalog: fonts, sizes, colors, spacing for every heading/body style
  3. Table styles: border, header formatting, alternating colors
  4. Page layout: margins, orientation, page size
  5. Numbering conventions: table/figure prefixes, list styles
  6. Header/footer text
"""
from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any

log = logging.getLogger(__name__)


def _hex_from_color(color_obj) -> str | None:
    """Safely extract hex color from a python-docx color object."""
    try:
        if color_obj and color_obj.rgb:
            return str(color_obj.rgb).upper()
    except Exception:
        pass
    return None


def _pt_from_emu(emu_val) -> float | None:
    """Convert EMU to points (1 pt = 12700 EMU)."""
    if emu_val is None:
        return None
    try:
        return round(emu_val / 12700, 1)
    except Exception:
        return None


def _pt_from_twips(twips) -> float | None:
    """Convert twips to points (1 pt = 20 twips)."""
    if twips is None:
        return None
    try:
        return round(twips / 20, 1)
    except Exception:
        return None


class TemplateAnalyzer:
    """Deep analysis of a DOCX template for structure and style extraction."""

    def analyze(self, docx_path: Path) -> dict:
        """Full template analysis.

        Returns a dict with:
          sections, style_catalog, table_style, page_layout, numbering,
          headers_footers, has_cover_page
        """
        try:
            from docx import Document
            from docx.oxml.ns import qn
            doc = Document(str(docx_path))
        except Exception as exc:
            log.error("TemplateAnalyzer: could not open %s: %s", docx_path, exc)
            return self._empty_analysis()

        sections = self._extract_sections(doc)
        style_catalog = self._extract_style_catalog(doc)
        table_style = self._extract_table_style(doc)
        page_layout = self._extract_page_layout(doc)
        numbering = self._infer_numbering_conventions(doc)
        headers_footers = self._extract_headers_footers(doc)

        return {
            "sections": sections,
            "style_catalog": style_catalog,
            "table_style": table_style,
            "page_layout": page_layout,
            "numbering": numbering,
            "headers_footers": headers_footers,
            "has_cover_page": self._detect_cover_page(doc),
        }

    # ------------------------------------------------------------------
    # Section Structure
    # ------------------------------------------------------------------

    def _extract_sections(self, doc) -> list[dict]:
        """Extract ordered section headings with hierarchy."""
        sections: list[dict] = []
        current_h1: dict | None = None

        for para in doc.paragraphs:
            text = para.text.strip()
            style_name = para.style.name if para.style else ""

            if not style_name.startswith("Heading"):
                continue
            if not text:
                continue

            try:
                level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                level = 1

            section = {
                "heading_text": text,
                "heading_level": level,
                "style_name": style_name,
                "subsections": [],
            }

            if level == 1:
                sections.append(section)
                current_h1 = section
            elif level == 2 and current_h1:
                current_h1["subsections"].append(section)
            else:
                sections.append(section)

        return sections

    # ------------------------------------------------------------------
    # Style Catalog
    # ------------------------------------------------------------------

    def _extract_style_catalog(self, doc) -> dict:
        """Extract formatting for each named style in the document."""
        catalog: dict[str, dict] = {}

        for style in doc.styles:
            style_name = style.name
            # Only capture heading and body-relevant styles
            relevant = (
                style_name.startswith("Heading")
                or style_name in ("Normal", "Body Text", "Caption", "List Bullet",
                                  "List Number", "Table Grid", "Title", "Subtitle",
                                  "Intense Quote", "Quote", "No Spacing")
            )
            if not relevant:
                continue

            info: dict[str, Any] = {"style_name": style_name}

            try:
                font = style.font
                if font:
                    if font.name:
                        info["font_family"] = font.name
                    if font.size:
                        info["font_size_pt"] = round(font.size.pt, 1)
                    if font.bold is not None:
                        info["bold"] = font.bold
                    if font.italic is not None:
                        info["italic"] = font.italic
                    if font.underline is not None:
                        info["underline"] = font.underline
                    color = _hex_from_color(font.color)
                    if color:
                        info["color_hex"] = color

                pf = style.paragraph_format
                if pf:
                    if pf.alignment is not None:
                        info["alignment"] = str(pf.alignment).split(".")[-1].lower()
                    if pf.space_before is not None:
                        info["space_before_pt"] = _pt_from_emu(pf.space_before)
                    if pf.space_after is not None:
                        info["space_after_pt"] = _pt_from_emu(pf.space_after)
                    if pf.line_spacing is not None:
                        # line_spacing can be a Length (pt) or a float multiplier
                        ls = pf.line_spacing
                        try:
                            info["line_spacing"] = round(float(ls), 2)
                        except Exception:
                            pass

            except Exception as exc:
                log.debug("Style extraction failed for %s: %s", style_name, exc)

            catalog[style_name] = info

        return catalog

    # ------------------------------------------------------------------
    # Table Style
    # ------------------------------------------------------------------

    def _extract_table_style(self, doc) -> dict:
        """Infer table formatting from the first table in the document."""
        table_info: dict[str, Any] = {
            "style_name": "Table Grid",
            "has_header_row": True,
            "header_bg_hex": None,
            "header_font_color_hex": None,
            "header_bold": True,
            "alt_row_colors": [],
            "border_color_hex": None,
            "cell_padding_pt": 3.0,
        }

        if not doc.tables:
            return table_info

        try:
            table = doc.tables[0]
            if table.style:
                table_info["style_name"] = table.style.name

            # Inspect header row (first row) formatting
            if table.rows:
                first_row = table.rows[0]
                for cell in first_row.cells:
                    # Try to get fill color from XML
                    try:
                        from docx.oxml.ns import qn
                        tc = cell._tc
                        tcPr = tc.find(qn("w:tcPr"))
                        if tcPr is not None:
                            shd = tcPr.find(qn("w:shd"))
                            if shd is not None:
                                fill = shd.get(qn("w:fill"))
                                if fill and fill != "auto" and len(fill) == 6:
                                    table_info["header_bg_hex"] = fill.upper()
                    except Exception:
                        pass

                    # Font color
                    for para in cell.paragraphs:
                        for run in para.runs:
                            color = _hex_from_color(run.font.color)
                            if color:
                                table_info["header_font_color_hex"] = color
                            if run.font.bold:
                                table_info["header_bold"] = True
                    break

            # Inspect second row for alternating color
            if len(table.rows) > 1:
                second_row = table.rows[1]
                for cell in second_row.cells:
                    try:
                        from docx.oxml.ns import qn
                        tc = cell._tc
                        tcPr = tc.find(qn("w:tcPr"))
                        if tcPr is not None:
                            shd = tcPr.find(qn("w:shd"))
                            if shd is not None:
                                fill = shd.get(qn("w:fill"))
                                if fill and fill != "auto" and len(fill) == 6:
                                    table_info["alt_row_colors"].append(fill.upper())
                    except Exception:
                        pass
                    break

        except Exception as exc:
            log.debug("Table style extraction failed: %s", exc)

        return table_info

    # ------------------------------------------------------------------
    # Page Layout
    # ------------------------------------------------------------------

    def _extract_page_layout(self, doc) -> dict:
        """Extract page dimensions and margins from the first section."""
        layout: dict[str, Any] = {
            "orientation": "portrait",
            "page_width_pt": 612.0,
            "page_height_pt": 792.0,
            "margins": {"top": 72.0, "bottom": 72.0, "left": 72.0, "right": 72.0},
        }

        try:
            if not doc.sections:
                return layout

            sect = doc.sections[0]
            w_pt = _pt_from_emu(sect.page_width)
            h_pt = _pt_from_emu(sect.page_height)
            if w_pt and h_pt:
                layout["page_width_pt"] = w_pt
                layout["page_height_pt"] = h_pt
                layout["orientation"] = "landscape" if w_pt > h_pt else "portrait"

            margins: dict[str, float | None] = {
                "top": _pt_from_emu(sect.top_margin),
                "bottom": _pt_from_emu(sect.bottom_margin),
                "left": _pt_from_emu(sect.left_margin),
                "right": _pt_from_emu(sect.right_margin),
            }
            layout["margins"] = {k: v for k, v in margins.items() if v is not None}

        except Exception as exc:
            log.debug("Page layout extraction failed: %s", exc)

        return layout

    # ------------------------------------------------------------------
    # Numbering / Conventions
    # ------------------------------------------------------------------

    def _infer_numbering_conventions(self, doc) -> dict:
        """Detect numbering prefixes for tables and figures from paragraph text."""
        conventions: dict[str, Any] = {
            "table_prefix": "Table",
            "figure_prefix": "Figure",
            "section_numbered": False,
        }

        table_pattern = re.compile(r"^(Table|Tabel|Tab\.?)\s*\d+", re.IGNORECASE)
        figure_pattern = re.compile(r"^(Figure|Fig\.?)\s*\d+", re.IGNORECASE)

        for para in doc.paragraphs:
            text = para.text.strip()
            if table_pattern.match(text):
                prefix = table_pattern.match(text).group(1)
                conventions["table_prefix"] = prefix.rstrip(".")
            if figure_pattern.match(text):
                prefix = figure_pattern.match(text).group(1)
                conventions["figure_prefix"] = prefix.rstrip(".")

        # Check if Heading 1 text starts with a number (e.g. "1. Introduction")
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            if style_name == "Heading 1":
                text = para.text.strip()
                if re.match(r"^\d+[\.\)]\s", text):
                    conventions["section_numbered"] = True
                break

        return conventions

    # ------------------------------------------------------------------
    # Headers / Footers
    # ------------------------------------------------------------------

    def _extract_headers_footers(self, doc) -> dict:
        """Extract text from page headers and footers."""
        result: dict[str, str | None] = {"header": None, "footer": None}

        try:
            if doc.sections:
                sect = doc.sections[0]
                header = sect.header
                if header and not header.is_linked_to_previous:
                    header_text = " ".join(p.text for p in header.paragraphs if p.text.strip())
                    if header_text:
                        result["header"] = header_text.strip()

                footer = sect.footer
                if footer and not footer.is_linked_to_previous:
                    footer_text = " ".join(p.text for p in footer.paragraphs if p.text.strip())
                    if footer_text:
                        result["footer"] = footer_text.strip()

        except Exception as exc:
            log.debug("Header/footer extraction failed: %s", exc)

        return result

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _detect_cover_page(self, doc) -> bool:
        """Heuristic: first paragraph has a 'Title' style → likely a cover page."""
        for para in doc.paragraphs:
            if para.text.strip():
                style_name = para.style.name if para.style else ""
                return "title" in style_name.lower()
        return False

    def _empty_analysis(self) -> dict:
        return {
            "sections": [],
            "style_catalog": {},
            "table_style": {},
            "page_layout": {},
            "numbering": {"table_prefix": "Table", "figure_prefix": "Figure"},
            "headers_footers": {},
            "has_cover_page": False,
        }
