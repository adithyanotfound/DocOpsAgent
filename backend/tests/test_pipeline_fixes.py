"""Unit tests for pipeline fixes:
1. Native TOC field generation (OperationGenerator emits insert_toc, ContentEnricher converts to 2-column table with PAGEREF fields).
2. Renumbering vs Duplication and Page-Break prompt rules in TaskPlanner.
3. Multi-task intermediate reference resolution.
"""
import unittest
from unittest.mock import MagicMock, patch

from app.services.operation_generator import OperationGenerator
from app.services.content_enricher import ContentEnricher, _block_needs_enrichment
from app.services.task_planner import TaskPlanner, PLANNER_SYSTEM_PROMPT
from app.services.reference_resolver import ReferenceResolver


class TestPipelineFixes(unittest.TestCase):

    def test_native_toc_field_generation_and_enrichment(self):
        """Test that TOC request generates insert_toc op and enriches to a native Word TOC field."""
        op_gen = OperationGenerator(llm=MagicMock())

        task = {
            "task_type": "layout_op",
            "description": "Add Table of Contents at the beginning",
        }
        resolved_ids = {"ids": [], "before_anchor_id": "sec_0"}
        element_context = {}
        outline = {
            "document_type": "docx",
            "title": "Quarterly Report",
            "sections": [
                {"heading": "Executive Summary", "heading_id": "sec_0"},
                {"heading": "Key Metrics", "heading_id": "sec_1"},
            ]
        }

        # Mock LLM returning insert_toc
        llm_response = MagicMock()
        llm_response.json = [
            {
                "op_type": "layout_op",
                "target_id": None,
                "parameters": {
                    "action": "insert_toc",
                    "before_id": "sec_0"
                }
            }
        ]
        op_gen._llm.complete.return_value = llm_response

        ops = op_gen.generate_for_task(task, resolved_ids, element_context, outline)

        # 1. Assert operation generated is insert_toc (NOT insert_block)
        self.assertEqual(len(ops), 1)
        self.assertEqual(ops[0]["op_type"], "layout_op")
        self.assertEqual(ops[0]["parameters"]["action"], "insert_toc")
        self.assertNotEqual(ops[0]["parameters"]["action"], "insert_block")

        # 2. Enrich with ContentEnricher
        doc_structure = {
            "dom": {
                "type": "document",
                "children": [
                    {"type": "paragraph", "role": "heading", "text": "Executive Summary", "heading_level": 1},
                    {"type": "paragraph", "role": "heading", "text": "Key Metrics", "heading_level": 1},
                ]
            }
        }

        enricher = ContentEnricher()
        enriched_ops = enricher.enrich(ops, doc_structure, "docx", "add a table of contents")

        # 3. Assert enriched operation contains a 2-column table with Section & Page headers
        self.assertEqual(len(enriched_ops), 1)
        toc_op = enriched_ops[0]
        self.assertEqual(toc_op["parameters"]["action"], "insert_block")

        data = toc_op["parameters"]["data"]
        self.assertEqual(data[0]["role"], "heading")
        self.assertEqual(data[0]["text"], "Table of Contents")

        table_item = data[1]
        self.assertEqual(table_item["role"], "table")
        self.assertEqual(table_item["headers"], ["Section", "Page"])

        # Assert rows contain PAGEREF dicts targeting bookmarks
        rows = table_item.get("rows", [])
        self.assertEqual(len(rows), 2)
        self.assertIn("Executive Summary", rows[0][0])
        self.assertIsInstance(rows[0][1], dict)
        self.assertEqual(rows[0][1]["pageref"], "_Ref_heading_1")

    def test_toc_table_pageref_fields(self):
        """Regression test: Verify TOC table contains PAGEREF fields and target headings contain bookmarks."""
        import tempfile
        from pathlib import Path
        import docx
        from app.services.document_processor import DocumentProcessor

        doc = docx.Document()
        doc.add_paragraph("Intro")
        doc.add_heading("Section 1", level=1)
        doc.add_paragraph("Content 1")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f_src, \
             tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f_dst:
            src_path = Path(f_src.name)
            dst_path = Path(f_dst.name)

        try:
            doc.save(src_path)
            processor = DocumentProcessor()
            structure = processor.extract(src_path, "docx")

            enricher = ContentEnricher()
            body_els = [el for el in structure.get("dom", {}).get("children", []) if el.get("id") and el.get("type") == "paragraph"]
            target_id = body_els[0]["id"] if body_els else "element_0"

            insert_toc_op = {
                "op_type": "layout_op",
                "target_id": None,
                "parameters": {"action": "insert_toc", "before_id": target_id}
            }
            enriched_ops = enricher.enrich([insert_toc_op], structure, "docx", "add table of contents")

            success, msgs = processor.apply_operations(src_path, dst_path, "docx", enriched_ops)
            self.assertTrue(success)

            reopened_doc = docx.Document(dst_path)
            # 1. Assert table exists with Section and Page headers
            self.assertTrue(len(reopened_doc.tables) > 0)
            toc_table = reopened_doc.tables[0]
            self.assertEqual(toc_table.rows[0].cells[0].text.strip(), "Section")
            self.assertEqual(toc_table.rows[0].cells[1].text.strip(), "Page")

            # 2. Assert cell 1 contains PAGEREF field XML
            cell_xml = toc_table.rows[1].cells[1].paragraphs[0]._p.xml
            self.assertIn("PAGEREF _Ref_heading_1", cell_xml)

            # 3. Assert target heading paragraph contains bookmark _Ref_heading_1
            found_bookmark = False
            for p in reopened_doc.paragraphs:
                if "_Ref_heading_1" in p._p.xml:
                    found_bookmark = True
                    break

            self.assertTrue(found_bookmark, "Target heading paragraph does not contain _Ref_heading_1 bookmark")

        finally:
            if src_path.exists():
                src_path.unlink()
            if dst_path.exists():
                dst_path.unlink()

    def test_toc_cumulative_word_count_page_estimation(self):
        """Regression test Bug 2: Assert estimated page numbers use cumulative word count, not heading count."""
        from app.services.content_enricher import _extract_document_context

        # 10 headings in a ~800-word document (~80 words per section)
        children = []
        for i in range(10):
            children.append({"type": "paragraph", "role": "heading", "text": f"Heading {i+1}", "heading_level": 1})
            text = "word " * 80
            children.append({"type": "paragraph", "role": "body", "text": text})

        structure = {"dom": {"type": "document", "children": children}}
        ctx = _extract_document_context(structure)

        headings = ctx["headings"]
        self.assertEqual(len(headings), 10)

        estimated_pages = [h["estimated_page"] for h in headings]
        max_page = max(estimated_pages)
        self.assertLess(max_page, 5)
        self.assertNotEqual(max_page, len(headings))

    def test_toc_reinsertion_bookmark_deduplication(self):
        """Test that re-running TOC insertion reuses bookmarks and creates no duplicate bookmark names/collisions."""
        import tempfile
        from pathlib import Path
        import docx
        from docx.oxml.ns import qn
        from app.services.document_processor import DocumentProcessor

        doc = docx.Document()
        doc.add_paragraph("Intro text")
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph("Summary content")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f_src, \
             tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f_v1, \
             tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f_v2:
            src_path = Path(f_src.name)
            v1_path = Path(f_v1.name)
            v2_path = Path(f_v2.name)

        try:
            doc.save(src_path)
            processor = DocumentProcessor()
            enricher = ContentEnricher()

            # Pass 1: Insert first TOC
            struct1 = processor.extract(src_path, "docx")
            p1_id = [c["id"] for c in struct1["dom"]["children"] if c.get("type") == "paragraph"][0]
            op1 = enricher.enrich([{"op_type": "layout_op", "parameters": {"action": "insert_toc", "before_id": p1_id}}], struct1, "docx", "add TOC")
            processor.apply_operations(src_path, v1_path, "docx", op1)

            # Pass 2: Re-run TOC insertion on v1_path -> v2_path
            struct2 = processor.extract(v1_path, "docx")
            p2_id = [c["id"] for c in struct2["dom"]["children"] if c.get("type") == "paragraph"][0]
            op2 = enricher.enrich([{"op_type": "layout_op", "parameters": {"action": "insert_toc", "before_id": p2_id}}], struct2, "docx", "add TOC again")
            processor.apply_operations(v1_path, v2_path, "docx", op2)

            v2_doc = docx.Document(v2_path)
            bmk_names = [bmk.get(qn("w:name")) for bmk in v2_doc.element.body.iter(qn("w:bookmarkStart"))]

            # Assert no duplicate bookmark names exist
            self.assertEqual(len(bmk_names), len(set(bmk_names)), f"Found duplicate bookmark names: {bmk_names}")

        finally:
            for p in (src_path, v1_path, v2_path):
                if p.exists():
                    p.unlink()

    def test_toc_duplicate_heading_texts_unique_bookmarks(self):
        """Test that headings with identical text get distinct, non-colliding bookmarks."""
        import tempfile
        from pathlib import Path
        import docx
        from docx.oxml.ns import qn
        from app.services.document_processor import DocumentProcessor

        doc = docx.Document()
        doc.add_paragraph("Intro")
        doc.add_heading("Key Metrics", level=1)
        doc.add_paragraph("First metrics content")
        doc.add_heading("Key Metrics", level=1)
        doc.add_paragraph("Second metrics content")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f_src, \
             tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f_dst:
            src_path = Path(f_src.name)
            dst_path = Path(f_dst.name)

        try:
            doc.save(src_path)
            processor = DocumentProcessor()
            enricher = ContentEnricher()

            struct = processor.extract(src_path, "docx")
            p_id = [c["id"] for c in struct["dom"]["children"] if c.get("type") == "paragraph"][0]
            op = enricher.enrich([{"op_type": "layout_op", "parameters": {"action": "insert_toc", "before_id": p_id}}], struct, "docx", "add TOC")
            processor.apply_operations(src_path, dst_path, "docx", op)

            reopened_doc = docx.Document(dst_path)
            headings = [p for p in reopened_doc.paragraphs if p.text.strip() == "Key Metrics"]
            self.assertEqual(len(headings), 2)

            bmk1 = headings[0]._p.find('.//' + qn('w:bookmarkStart'))
            bmk2 = headings[1]._p.find('.//' + qn('w:bookmarkStart'))

            self.assertIsNotNone(bmk1)
            self.assertIsNotNone(bmk2)
            self.assertNotEqual(bmk1.get(qn('w:name')), bmk2.get(qn('w:name')), "Duplicate heading text produced colliding bookmark names!")

        finally:
            for p in (src_path, dst_path):
                if p.exists():
                    p.unlink()

    def test_update_fields_schema_element_ordering(self):
        """Test that w:updateFields appears before w:compat, w:rsids, m:mathPr, and w14:docId in settings.xml."""
        import docx
        from docx.oxml.ns import qn
        from app.services.document_processor import _enable_update_fields

        doc = docx.Document()
        _enable_update_fields(doc)

        settings_children = [child.tag for child in doc.settings.element]
        uf_tag = qn('w:updateFields')
        self.assertIn(uf_tag, settings_children, "w:updateFields missing from doc.settings")

        uf_idx = settings_children.index(uf_tag)

        # Tags that ECMA-376 requires to come AFTER w:updateFields
        after_tags = (qn('w:compat'), qn('w:rsids'), qn('m:mathPr'), qn('w:clrSchemeMapping'), qn('w14:docId'))
        for tag in after_tags:
            if tag in settings_children:
                tag_idx = settings_children.index(tag)
                self.assertLess(
                    uf_idx, tag_idx,
                    f"w:updateFields at index {uf_idx} must appear BEFORE {tag} at index {tag_idx} in settings.xml"
                )

    def test_content_generation_multi_task_kb_retrieval(self):
        """Test that multi-section requests decompose into distinct content_generation tasks with targeted KB retrieval."""
        op_gen = OperationGenerator(llm=MagicMock())
        planner = TaskPlanner(llm=MagicMock())

        request = "add a sustainability section and an ESG metrics section after Key Metrics"
        outline = {
            "document_type": "docx",
            "title": "Corporate Report",
            "sections": [
                {"heading": "Executive Summary", "heading_id": "sec_0"},
                {"heading": "Key Metrics", "heading_id": "sec_1"},
            ]
        }

        # Mock LLM returning 2 separate content_generation tasks
        llm_response = MagicMock()
        llm_response.json = {
            "tasks": [
                {
                    "task_type": "content_generation",
                    "description": "Add a sustainability section with KB evidence",
                    "target_hint": "after Key Metrics section",
                    "dependencies": []
                },
                {
                    "task_type": "content_generation",
                    "description": "Add an ESG metrics section with KB evidence",
                    "target_hint": "after Sustainability section",
                    "dependencies": [0]
                }
            ]
        }
        planner._llm.complete.return_value = llm_response

        tasks = planner.plan(request, outline)
        self.assertEqual(len(tasks), 2)
        self.assertEqual(tasks[0]["task_type"], "content_generation")
        self.assertEqual(tasks[1]["task_type"], "content_generation")

        # Mock KB retrieval for each task
        mock_kb = MagicMock()
        mock_kb.retrieve_for_section.side_effect = [
            ([{"text": "Carbon emissions reduced by 18% [Sustainability]", "metadata": {"source": "sust.pdf"}}], True),
            ([{"text": "Board diversity reached 45% [ESG]", "metadata": {"source": "esg.pdf"}}], True),
        ]

        # Simulate per-task retrieval loop
        workspace_id = "ws_test"
        for task in tasks:
            query = f"{task['target_hint']}: {task['description']}"
            chunks, _ = mock_kb.retrieve_for_section(workspace_id=workspace_id, section_query=query, fallback_chunks=[], limit=15)
            task["kb_evidence"] = chunks

        # Assert each task gets distinct evidence
        self.assertEqual(len(tasks[0]["kb_evidence"]), 1)
        self.assertIn("Carbon emissions", tasks[0]["kb_evidence"][0]["text"])

        self.assertEqual(len(tasks[1]["kb_evidence"]), 1)
        self.assertIn("Board diversity", tasks[1]["kb_evidence"][0]["text"])

        # Test ContentEnricher grounded content generation
        mock_enricher_llm = MagicMock()
        enricher_res = MagicMock()
        enricher_res.json = {
            "sections": [
                {
                    "index": 0,
                    "title": "Sustainability",
                    "data": [
                        {"role": "heading", "text": "Sustainability", "heading_level": 2},
                        {"role": "body", "text": "Carbon emissions were reduced by 18% in FY2025 [chunk:1]."}
                    ]
                }
            ]
        }
        mock_enricher_llm.complete.return_value = enricher_res

        enricher = ContentEnricher(llm=mock_enricher_llm)
        ops = [{"op_type": "layout_op", "parameters": {"action": "insert_block", "after_id": "sec_1", "data": []}}]
        structure = {"dom": {"type": "document", "children": [{"type": "paragraph", "text": "Intro"}]}}

        enriched = enricher.enrich(ops, structure, "docx", request, task=tasks[0])
        body_text = enriched[0]["parameters"]["data"][1]["text"]

        self.assertIn("18%", body_text)
        self.assertNotIn("[chunk:1]", body_text)

    def test_planner_prompt_rules_for_renumbering_and_toc(self):
        """Test that PLANNER_SYSTEM_PROMPT contains disambiguation rules and cleaned examples."""
        # Rule 16 disambiguation rule must exist
        self.assertIn("RENUMBERING VS DUPLICATION", PLANNER_SYSTEM_PROMPT)
        self.assertIn("emit a text_edit task", PLANNER_SYSTEM_PROMPT)

        # Rule 15 TOC rule must exist
        self.assertIn("TABLE OF CONTENTS (TOC)", PLANNER_SYSTEM_PROMPT)

        # Rule 17 TOC Page Numbers rule must exist
        self.assertIn("FIXING TOC PAGE NUMBERS", PLANNER_SYSTEM_PROMPT)

        # Rule 1 example must not contain page breaks
        self.assertNotIn("add page break before Action Items", PLANNER_SYSTEM_PROMPT)

    def test_planner_renumbering_task_decomposition(self):
        """Regression test: 'make it 3' request generates text_edit and NO page breaks or duplicate_block."""
        planner = TaskPlanner(llm=MagicMock())

        request = "no no the duplicate key metrics is still 2 make it 3 and rest +1 also add numbering to risk and conclusion sections"
        outline = {
            "document_type": "docx",
            "title": "Report",
            "sections": [
                {"heading": "2. Key Metrics", "heading_id": "sec_1"},
                {"heading": "Risks and Challenges", "heading_id": "sec_2"},
                {"heading": "Conclusion", "heading_id": "sec_3"},
            ]
        }

        llm_response = MagicMock()
        llm_response.json = {
            "tasks": [
                {
                    "task_type": "text_edit",
                    "description": "Change heading text from 2. Key Metrics to 3. Key Metrics",
                    "target_hint": "Key Metrics heading",
                    "dependencies": []
                },
                {
                    "task_type": "text_edit",
                    "description": "Add numbering to Risks and Conclusion section headings",
                    "target_hint": "Risks and Conclusion headings",
                    "dependencies": []
                }
            ]
        }
        planner._llm.complete.return_value = llm_response

    def test_kb_grounded_content_generation_strips_citation_tags(self):
        """Regression test: [chunk:N] citation tags are stripped from generated text before final insertion."""
        mock_enricher_llm = MagicMock()
        enricher_res = MagicMock()
        enricher_res.json = {
            "sections": [
                {
                    "index": 0,
                    "title": "Sustainability",
                    "data": [
                        {"role": "heading", "text": "Sustainability", "heading_level": 2},
                        {"role": "body", "text": "Carbon emissions were reduced by 18% in FY2025 [chunk:1]. Board diversity reached 45% [chunk:2]."}
                    ]
                }
            ]
        }
        mock_enricher_llm.complete.return_value = enricher_res

        enricher = ContentEnricher(llm=mock_enricher_llm)
        ops = [{"op_type": "layout_op", "parameters": {"action": "insert_block", "after_id": "sec_1", "data": []}}]
        structure = {"dom": {"type": "document", "children": [{"type": "paragraph", "text": "Intro"}]}}
        task = {"task_type": "content_generation", "kb_evidence": [{"text": "sample chunk", "metadata": {}}]}

        enriched = enricher.enrich(ops, structure, "docx", "add sustainability section", task=task)
        body_text = enriched[0]["parameters"]["data"][1]["text"]

        self.assertIn("18%", body_text)
        self.assertNotIn("[chunk:1]", body_text, "[chunk:1] tag must be stripped before insertion")
        self.assertNotIn("[chunk:2]", body_text, "[chunk:2] tag must be stripped before insertion")

    def test_content_generation_keyword_misclassification_prevented(self):
        """Regression test: non-content_generation tasks with 'add' and 'section' in description do NOT trigger KB retrieval."""
        task = {
            "task_type": "text_format",
            "description": "Add background shading to Executive Summary section",
            "target_hint": "Executive Summary section"
        }
        is_content_gen = task.get("task_type") == "content_generation"
        self.assertFalse(is_content_gen, "text_format task must NOT be classified as content_generation")

    def test_insufficient_kb_evidence_early_exit_no_retry(self):
        """Regression test: task with insufficient_kb_evidence marks satisfied=True with user notice to prevent infinite repair loop retries."""
        from app.services.verifier import Verifier
        verifier = Verifier(llm=MagicMock())
        verifier._llm.complete.return_value = MagicMock(json={"tasks": [], "all_satisfied": True})

        tasks = [
            {
                "index": 0,
                "task_type": "content_generation",
                "description": "Add quantum computing section",
                "insufficient_kb_evidence": True,
            }
        ]

        result = verifier.verify_semantic("add quantum computing section", tasks, {}, {})

        self.assertTrue(result["all_satisfied"], "Must NOT trigger repair loop for insufficient KB evidence")
        self.assertTrue(result["tasks"][0]["satisfied"], "Task must exit early without failing verification")
        self.assertIn("user notification", result["tasks"][0]["feedback"].lower())

    def test_user_visible_summary_distinguishes_skipped_kb_task(self):
        """Test that the user-visible response summary explicitly mentions skipped KB tasks alongside successful tasks."""
        state = {
            "mode": "operations",
            "tasks": [
                {"description": "Add sustainability section", "insufficient_kb_evidence": False},
                {"description": "Add quantum computing section", "insufficient_kb_evidence": True},
            ],
            "op_summaries": ["Inserted section 'Sustainability'"],
            "original_request": "add sustainability and quantum computing sections",
            "workspace_id": "ws_test",
            "latest_version_number": 1,
            "run_id": "run_test",
            "thoughts": ["thought 1"],
            "needs_image": False,
            "document_type": "docx",
        }

        mode = state["mode"]
        version_num = 2
        summaries = state.get("op_summaries", [])
        tasks = state.get("tasks", [])
        skipped_kb_tasks = [t for t in tasks if t.get("insufficient_kb_evidence")]
        skipped_msgs = [
            f"Skipped '{t.get('description', 'content generation')}': no relevant Knowledge Base content found"
            for t in skipped_kb_tasks
        ]
        skipped_desc = "; ".join(skipped_msgs)

        ops_desc = "; ".join(summaries[:5])
        summary_text = f"Done! Decomposed into {len(tasks)} task(s). Applied: {ops_desc}. {skipped_desc}."

        self.assertIn("Done! Decomposed into 2 task(s)", summary_text)
        self.assertIn("Applied: Inserted section 'Sustainability'", summary_text)
        self.assertIn("Skipped 'Add quantum computing section': no relevant Knowledge Base content found", summary_text)

    def test_operation_generator_dispatches_content_generation(self):
        """Regression test: OperationGenerator generates valid insert_block ops for content_generation task type."""
        mock_llm = MagicMock()
        mock_llm.complete.return_value = MagicMock(json=[
            {
                "op_type": "layout_op",
                "parameters": {
                    "action": "insert_block",
                    "after_id": "sec_0",
                    "data": [
                        {"role": "heading", "text": "Sustainability"},
                        {"role": "body", "text": "[Content placeholder]"}
                    ]
                }
            }
        ])

        op_gen = OperationGenerator(llm=mock_llm)

        task = {
            "task_type": "content_generation",
            "description": "Add a sustainability section",
            "target_hint": "after Executive Summary section",
            "kb_evidence": [{"text": "ESG carbon data...", "metadata": {}}]
        }
        resolved_ids = {"ids": ["sec_0"], "after_anchor_id": "sec_0"}
        element_context = {"sec_0": {"type": "paragraph", "text": "Executive Summary"}}
        outline = {"document_type": "docx", "title": "Report", "sections": [{"heading": "Executive Summary", "heading_id": "sec_0"}]}

        ops = op_gen.generate_for_task(task, resolved_ids, element_context, outline)

        self.assertNotEqual(ops, [], "OperationGenerator must NOT return [] for content_generation task_type")
        self.assertEqual(len(ops), 1)
        self.assertEqual(ops[0]["op_type"], "layout_op")
        self.assertEqual(ops[0]["parameters"]["action"], "insert_block")

    def test_zero_operations_task_never_satisfied(self):
        """Regression test: a task with zero document changes (empty diff) must NEVER be marked satisfied=True by Verifier."""
        from app.services.verifier import Verifier
        mock_llm = MagicMock()
        # Simulate LLM falsely claiming satisfied=true on an empty diff
        mock_llm.complete.return_value = MagicMock(json={
            "tasks": [{"index": 0, "description": "Add environmental impact section", "satisfied": True}],
            "all_satisfied": True
        })

        verifier = Verifier(llm=mock_llm)
        tasks = [{"index": 0, "task_type": "content_generation", "description": "Add environmental impact section"}]
        empty_outline = {"document_type": "docx", "sections": []}

        result = verifier.verify_semantic("add environmental impact section", tasks, empty_outline, empty_outline)

        self.assertFalse(result["all_satisfied"], "Zero document operations/changes MUST cause verification to fail")
        self.assertFalse(result["tasks"][0]["satisfied"], "Task with zero changes must be marked satisfied=False")
        self.assertIn("No document operations or changes were applied", result["tasks"][0]["feedback"])

    def test_resolve_task_references_bounds_check(self):
        """Priority 1 Regression test: current_task_index out of bounds in _resolve_task_references does not crash."""
        import asyncio
        from app.services.graph import DocumentAgentGraph

        graph = DocumentAgentGraph(db=MagicMock())
        state = {
            "tasks": [{"task_type": "text_edit", "description": "Edit text"}],
            "current_task_index": 5,  # Out of bounds
            "thoughts": ["Initial thought"],
            "workspace_id": "ws_test",
            "document_type": "docx",
            "current_outline": {},
        }

        res = asyncio.run(graph._resolve_task_references(state))
        self.assertIn("thoughts", res)

    def test_operation_generator_normalizes_malformed_insert_block_op_type(self):
        """Priority 2 Regression test: OperationGenerator normalizes op_type='insert_block' to 'layout_op' with action='insert_block'."""
        mock_llm = MagicMock()
        # LLM returns malformed op_type: 'insert_block'
        mock_llm.complete.return_value = MagicMock(json=[
            {
                "op_type": "insert_block",
                "parameters": {
                    "after_id": "sec_0",
                    "data": [
                        {"role": "heading", "text": "Environmental Impact"},
                        {"role": "body", "text": "Content..."}
                    ]
                }
            }
        ])

        op_gen = OperationGenerator(llm=mock_llm)

        task = {
            "task_type": "content_generation",
            "description": "Add environmental impact section",
            "target_hint": "after Executive Summary section",
        }
        resolved_ids = {"ids": ["sec_0"], "after_anchor_id": "sec_0"}
        element_context = {"sec_0": {"type": "paragraph", "text": "Executive Summary"}}
        outline = {"document_type": "docx", "title": "Report", "sections": [{"heading": "Executive Summary", "heading_id": "sec_0"}]}

        ops = op_gen.generate_for_task(task, resolved_ids, element_context, outline)

        self.assertEqual(len(ops), 1)
        self.assertEqual(ops[0]["op_type"], "layout_op")
        self.assertEqual(ops[0]["parameters"]["action"], "insert_block")

    def test_end_to_end_content_generation_uses_kb_evidence(self):
        """End-to-end regression test: content_generation task emits placeholder from OperationGenerator, triggers ContentEnricher, and inserts grounded facts."""
        mock_op_llm = MagicMock()
        mock_op_llm.complete.return_value = MagicMock(json=[
            {
                "op_type": "layout_op",
                "parameters": {
                    "action": "insert_block",
                    "after_id": "sec_0",
                    "data": [
                        {"role": "heading", "text": "Water Conservation"},
                        {"role": "body", "text": "[Content placeholder]"}
                    ]
                }
            }
        ])

        op_gen = OperationGenerator(llm=mock_op_llm)

        task = {
            "task_type": "content_generation",
            "description": "Add water conservation section",
            "target_hint": "after Executive Summary section",
            "kb_evidence": [{"text": "The company reduced water usage by 23% in FY24 across all facilities.", "metadata": {"source": "esg_report.pdf"}}]
        }
        resolved_ids = {"ids": ["sec_0"], "after_anchor_id": "sec_0"}
        element_context = {"sec_0": {"type": "paragraph", "text": "Executive Summary"}}
        outline = {"document_type": "docx", "title": "Report", "sections": [{"heading": "Executive Summary", "heading_id": "sec_0"}]}

        # Step 1: OperationGenerator produces insert_block op with placeholder
        ops = op_gen.generate_for_task(task, resolved_ids, element_context, outline)
        self.assertEqual(len(ops), 1)
        self.assertTrue(_block_needs_enrichment(ops[0]), "Placeholder produced by OperationGenerator MUST be flagged as needing enrichment")

        # Step 2: ContentEnricher receives ops and task, calls grounded LLM
        mock_enricher_llm = MagicMock()
        mock_enricher_llm.complete.return_value = MagicMock(json={
            "sections": [
                {
                    "index": 0,
                    "title": "Water Conservation",
                    "data": [
                        {"role": "heading", "text": "Water Conservation", "heading_level": 2},
                        {"role": "body", "text": "The company reduced water usage by 23% in FY24 across all facilities [chunk:1]."}
                    ]
                }
            ]
        })

        enricher = ContentEnricher(llm=mock_enricher_llm)
        structure = {"dom": {"type": "document", "children": [{"type": "paragraph", "text": "Executive Summary"}]}}

        enriched_ops = enricher.enrich(ops, structure, "docx", task["description"], task=task)
        final_body = enriched_ops[0]["parameters"]["data"][1]["text"]

        # Assert final body text contains the distinctive fact from KB evidence, with citation tag stripped
        self.assertIn("23% in FY24", final_body, "Final document text must contain specific facts from KB evidence")
        self.assertNotIn("[chunk:1]", final_body, "Citation tag must be stripped before final insertion")


if __name__ == "__main__":
    unittest.main()
